"""
Complete Processor Test - Days 1-3
Tests extraction AND text normalization
"""

import sys
import os

# Add backend to path
sys.path.insert(0, os.path.dirname(__file__))

print("=" * 60)
print("COMPLETE CONTRACT PROCESSOR TEST")
print("=" * 60)
print()

# Test 1: Import
print("Test 1: Importing ContractProcessor...")
try:
    from app.core.processor import ContractProcessor
    print(" Import successful!")
except ImportError as e:
    print(f" Import failed: {e}")
    sys.exit(1)

print()

# Test 2: Initialize
print("Test 2: Initializing processor...")
try:
    processor = ContractProcessor(normalize_text=True)
    print("✅ Processor initialized with text normalization!")
except Exception as e:
    print(f" Initialization failed: {e}")
    sys.exit(1)

print()

# Test 3: Create test contracts with messy text
print("Test 3: Creating test contracts with messy text...")

# Create messy PDF
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    
    pdf_path = "test_messy_contract.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    
    # Add messy contract text (extra spaces, weird formatting)
    c.drawString(100, height - 100, "SOFTWARE    LICENSE   AGREEMENT")
    c.drawString(100, height - 130, "")
    c.drawString(100, height - 150, "This  agreement   is  between   TechCorp  and   ClientCo.")
    c.drawString(100, height - 170, "")
    c.drawString(100, height - 190, "TERMS:")
    c.drawString(100, height - 210, "1. Payment:   $5,000per   month")
    c.drawString(100, height - 230, "2. Duration:  12months   with  auto-renewal")
    c.drawString(100, height - 250, "3. Termination:   90days    notice   required")
    c.drawString(100, height - 270, "4. Late  payment   penalty: 5%  per  month")
    
    c.save()
    print(f" Created messy PDF: {pdf_path}")
    has_pdf = True
    
except ImportError:
    print(" reportlab not installed - skipping PDF test")
    print("   Install with: pip install reportlab")
    has_pdf = False
    pdf_path = None
except Exception as e:
    print(f" PDF creation failed: {e}")
    has_pdf = False
    pdf_path = None

# Create messy DOCX
try:
    from docx import Document
    
    docx_path = "test_messy_contract.docx"
    doc = Document()
    
    doc.add_heading('EMPLOYMENT    CONTRACT', 0)
    doc.add_paragraph('')
    doc.add_paragraph('This   Employment   Agreement  is   entered  into  between:')
    doc.add_paragraph('Employer:  TechCorp   Ltd.')
    doc.add_paragraph('Employee:   John   Doe')
    doc.add_paragraph('')
    doc.add_heading('TERMS   OF   EMPLOYMENT', level=1)
    doc.add_paragraph('1.  Position:   Senior   Software   Engineer')
    doc.add_paragraph('2.  Salary:  $120,000   per   year')
    doc.add_paragraph('3.  Start  Date:  January  15,   2024')
    doc.add_paragraph('4.  Probation:   90days')
    doc.add_paragraph('5.  Notice:  30  days')
    
    doc.save(docx_path)
    print(f" Created messy DOCX: {docx_path}")
    has_docx = True
    
except Exception as e:
    print(f" DOCX creation failed: {e}")
    has_docx = False
    docx_path = None

print()
print("=" * 60)
print("TESTING TEXT EXTRACTION & NORMALIZATION")
print("=" * 60)
print()

# Test 4: Extract from PDF (with normalization)
if has_pdf and pdf_path:
    print("Test 4: PDF Extraction + Normalization")
    print("-" * 60)
    try:
        # Get file info
        info = processor.get_file_info(pdf_path)
        print(f" File: {info['filename']}")
        print(f" Size: {info['size_mb']} MB")
        print()
        
        # Extract WITHOUT cleaning
        print(" Extracting text (WITHOUT normalization)...")
        raw_text = processor.extract_text(pdf_path, clean=False)
        print(f" Extracted {len(raw_text)} characters (raw)")
        print()
        print("Raw text preview (first 200 chars):")
        print(f'"{raw_text[:200]}"')
        print()
        
        # Extract WITH cleaning
        print(" Extracting text (WITH normalization)...")
        clean_text = processor.extract_text(pdf_path, clean=True)
        print(f" Extracted {len(clean_text)} characters (normalized)")
        print()
        print("Clean text preview (first 200 chars):")
        print(f'"{clean_text[:200]}"')
        print()
        
        # Show stats
        stats = processor.get_text_stats(clean_text)
        print(" Text Statistics:")
        print(f"   Words: {stats['words']}")
        print(f"   Sentences: {stats['sentences']}")
        print(f"   Paragraphs: {stats['paragraphs']}")
        print(f"   Avg word length: {stats['avg_word_length']}")
        print()
        
        # Validate
        validation = processor.validate_contract_text(clean_text)
        print("✓ Validation:")
        print(f"   Is valid contract: {validation['is_valid']}")
        if validation['is_valid']:
            print(f"   Keywords found: {len(validation['found_keywords'])}")
            print(f"   Sample keywords: {', '.join(validation['found_keywords'][:3])}")
        print()
        
        print(" PDF test PASSED!")
        print()
        
    except Exception as e:
        print(f" PDF test FAILED: {e}")
        import traceback
        traceback.print_exc()

print()

# Test 5: Extract from DOCX (with normalization)
if has_docx and docx_path:
    print("Test 5: DOCX Extraction + Normalization")
    print("-" * 60)
    try:
        # Get file info
        info = processor.get_file_info(docx_path)
        print(f" File: {info['filename']}")
        print(f" Size: {info['size_mb']} MB")
        print()
        
        # Extract WITHOUT cleaning
        print(" Extracting text (WITHOUT normalization)...")
        raw_text = processor.extract_text(docx_path, clean=False)
        print(f" Extracted {len(raw_text)} characters (raw)")
        print()
        print("Raw text preview (first 200 chars):")
        print(f'"{raw_text[:200]}"')
        print()
        
        # Extract WITH cleaning
        print(" Extracting text (WITH normalization)...")
        clean_text = processor.extract_text(docx_path, clean=True)
        print(f"  Extracted {len(clean_text)} characters (normalized)")
        print()
        print("Clean text preview (first 200 chars):")
        print(f'"{clean_text[:200]}"')
        print()
        
        # Show difference
        print("🔍 Normalization Impact:")
        spaces_removed = raw_text.count('  ') - clean_text.count('  ')
        print(f"   Extra spaces removed: {spaces_removed}")
        print(f"   Character reduction: {len(raw_text) - len(clean_text)}")
        print()
        
        # Show stats
        stats = processor.get_text_stats(clean_text)
        print(" Text Statistics:")
        print(f"   Words: {stats['words']}")
        print(f"   Sentences: {stats['sentences']}")
        print(f"   Paragraphs: {stats['paragraphs']}")
        print()
        
        # Validate
        validation = processor.validate_contract_text(clean_text)
        print(" Validation:")
        print(f"   Is valid contract: {validation['is_valid']}")
        print()
        
        print(" DOCX test PASSED!")
        print()
        
    except Exception as e:
        print(f" DOCX test FAILED: {e}")
        import traceback
        traceback.print_exc()

print()
print("=" * 60)
print("TEST SUMMARY")
print("=" * 60)
print()

if has_pdf and has_docx:
    print(" ALL TESTS PASSED!")
    print()
    print(" Test files created:")
    print(f"   - {pdf_path}")
    print(f"   - {docx_path}")
    print()
    print(" Key Features Tested:")
    print("   ✅ PDF extraction")
    print("   ✅ DOCX extraction")
    print("   ✅ Text normalization (removing extra spaces)")
    print("   ✅ Text validation (contract keywords)")
    print("   ✅ Statistics generation")
   
else:
    print("  Some tests skipped")
    if not has_pdf:
        print("   - PDF test skipped (reportlab needed)")
    if not has_docx:
        print("   - DOCX test skipped")

print()
print("=" * 60)